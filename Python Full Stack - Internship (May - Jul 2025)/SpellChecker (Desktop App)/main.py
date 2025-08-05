import sys #Provides access to system-specific parameters and functions.
from grammar_Checker import check_grammar
from spell_Checker import check_spelling
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QVBoxLayout, QPushButton, QTextEdit, QLabel, QWidget, QListWidget, QSplitter, QMessageBox
from PyQt5.QtGui import QTextCharFormat, QColor, QSyntaxHighlighter, QPixmap, QPainter, QPen, QFont
from PyQt5.QtCore import Qt, QRect, QSize, QRegExp
from docx import Document

class Highlighter(QSyntaxHighlighter):
    def __init__(self, document):
        """
        Initializes the Highlighter object to highlight text with issues.

        Args:
            document (QTextDocument): The text document to be highlighted.
            issues (list): A list of dictionaries containing details about the issues.
        """
        super(Highlighter, self).__init__(document)
        self.issues = []
        self.setDocument(document)  # Ensure the highlighter is set with the document

    def set_issues(self, issues):
        self.issues = issues
        self.rehighlight()  # Apply highlighting immediately

    def highlightBlock(self, text):
        """
        Highlights the spelling errors in the block of text.
        
        :param text: The current block of text to be checked and highlighted.
        """
        # Define the format for highlighting spelling errors
        error_format = QTextCharFormat()
        error_format.setForeground(QColor("white"))
        error_format.setBackground(QColor("red"))

        # Highlight each word that is in the error_words list
        for issue in self.issues:
            word = issue['word'] if 'sentence' not in issue.keys() else issue['sentence']
            expression = QRegExp(f"\\b{word}\\b")  # Regular expression for whole words
            index = expression.indexIn(text)
            length = len(word)
            self.setFormat(index, length, error_format)
            
class CircularProgressBar(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.value = 0  # Current progress value (0 to 100)
        self.width = 100  # Size of the progress bar
        self.height = 100

    def setValue(self, value):
        self.value = value
        self.update()  # Trigger a repaint

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)

        # Define the rectangle for the progress bar
        rect = QRect(50, 50, self.width, self.height)

        # Draw the background circle
        painter.setPen(Qt.NoPen)
        painter.setBrush(QColor(0, 0, 0))  # Light gray color
        painter.drawEllipse(rect)

        # Draw the progress arc
        pen = QPen(QColor(0, 100, 255), 5)  # Set arc color and thickness
        painter.setPen(pen)
        startAngle = int(-90 * 16)  # Convert start angle to integer
        spanAngle = int(-self.value * 3.6 * 16)  # Convert span angle to integer
        painter.drawArc(rect, startAngle, spanAngle)

        # Draw the percentage text in the center
        painter.setPen(QColor(0, 170, 255))
        painter.setFont(QFont("Arial", 20, QFont.Bold))
        painter.drawText(rect, Qt.AlignCenter, f"{self.value}%")

    def sizeHint(self):
        return QSize(self.width, self.height)  # Return QSize object
       
class MainWindow(QMainWindow):  
    def __init__(self):
        """
        Initializes the main window of the application, setting up the layout,
        widgets, and connections to the check functions.
        """
        super().__init__()

        # create Window 
        self.createWindow("Spell and Grammar Checker","black","white")
        
        # Title and logo
        self.title_label = self.createProductTitle("SPELLGRA checker")
        self.logo_label = self.createLogo("logo.png")
        
        # Create a text editor widget for input
        self.text_edit = QTextEdit(self)
        
        # Create a list widget to display recommendations
        self.recommendations_box = QListWidget(self)
        
        # Buttons
        self.upload_btn = self.createButton("Upload File")
        self.check_grammar_button = self.createButton("Check Grammar")
        self.check_spelling_button = self.createButton("Check Spelling")
        
        # After click button
        self.upload_btn.clicked.connect(self.upload_file)   
        self.check_grammar_button.clicked.connect(self.run_grammar_checker)
        self.check_spelling_button.clicked.connect(self.run_spell_checker)
        
        # create progress bar
        self.progress_bar = CircularProgressBar(self)
        
        # create Layout
        logo_layout = self.createLayout(self.logo_label)
        title_layout = self.createLayout(self.title_label)
        text_layout = self.createLayout(self.text_edit, self.recommendations_box)
        button_layout = self.createLayout(self.upload_btn, self.check_grammar_button, self.check_spelling_button, self.progress_bar)
        
        top_layouts = [logo_layout, title_layout]
        bottom_layouts = [text_layout, button_layout]
        
        # Set windows split 
        horizontal_splitter_layout = []
        
        horizontal_splitter_layout.append(self.createWindowSplitter(Qt.Horizontal, top_layouts, 20, 80))
        horizontal_splitter_layout.append(self.createWindowSplitter(Qt.Horizontal, bottom_layouts, 70, 30))
        vertical_splitter_layout = self.createWindowSplitter(Qt.Vertical, horizontal_splitter_layout, 10, 90)
        
        self.main_layout = self.createLayout(vertical_splitter_layout)
               
        # Set the central widget
        self.setCentralWidget(self.main_layout)
        
    def createWindow(self,title,bgColor,fgColor):
        # Set window properties
        self.setWindowTitle(title)
        self.setGeometry(100, 100, 800, 600)
        self.setStyleSheet(f"background-color: {bgColor}; color: {fgColor};")
        
    def createProductTitle(self, titleName):
        title_label = QLabel(titleName)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("font-size: 24px; font-weight: bold; color: white;")
        return title_label
        
    def createLogo(self,logoName):
        logo_label = QLabel()
        logo_pixmap = QPixmap(f"resources/logo/{logoName}")  # Ensure you have this image file
        if not logo_pixmap.isNull():  # Check if the image is loaded successfully
            shrinked_pixmap = logo_pixmap.scaled(logo_pixmap.width() // 6, 
                                                 logo_pixmap.height() // 6, 
                                                 Qt.KeepAspectRatio, 
                                                 Qt.SmoothTransformation)
            logo_label.setPixmap(shrinked_pixmap)
            logo_label.setAlignment(Qt.AlignCenter)
        else:
            logo_label.setText("Logo not found")
            logo_label.setAlignment(Qt.AlignCenter)
        return logo_label
    
    def createButton(self, buttonName):
        button = QPushButton(buttonName,self)
        button.setStyleSheet("""
                                QPushButton {
                                    background-color: #4CAF50;  /* Green background */
                                    border: none;
                                    color: white;
                                    padding: 10px 20px;
                                    text-align: center;
                                    text-decoration: none;
                                    font-size: 16px;
                                    border-radius: 10px;  /* Rounded corners */
                                    transition: background-color 0.3s;
                                }
                                QPushButton:hover {
                                    background-color: #45a049;  /* Darker green on hover */
                                }
                            """)
        return button
        
    def createLayout(self,*widgets):
        widgetObj = QWidget()
        layout = QVBoxLayout()
        [layout.addWidget(widget) for widget in widgets]
        widgetObj.setLayout(layout)
        return widgetObj
        
    def createWindowSplitter(self, splitBy,widgetObjList,firstWidth,secondWidth):
        splitter = QSplitter(splitBy)
        splitter.setHandleWidth(1)  # Set the width of the splitter handle
        splitter.setStyleSheet("QSplitter::handle { background: grey; }")  # Style the splitter handle
        [splitter.addWidget(widgetObj) for widgetObj in widgetObjList]

        # Set the initial sizes to achieve a split
        initial_sizes = [int(self.width() *(firstWidth/100)), int(self.width() * (secondWidth/100))]
        splitter.setSizes(initial_sizes)
        return splitter
        
    def upload_file(self):
        # Open file dialog to select a file
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Open File", "", 
                                                   "Text Files (*.txt);;Word Files (*.docx);;All Files (*)", 
                                                   options=options)
        if file_name:
            try:
                if file_name.endswith('.txt'):
                    with open(file_name, 'r', encoding='utf-8') as file:
                        self.text_edit.setText(file.read())
                elif file_name.endswith('.docx'):
                    doc = Document(file_name)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    self.text_edit.setText('\n'.join(full_text))
                else:
                    QMessageBox.warning(self, "Unsupported File", "This file format is not supported.")
            except UnicodeDecodeError:
                QMessageBox.warning(self, "Error", "Cannot decode file with utf-8 encoding. Please choose a different encoding.")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to read file: {str(e)}")
        
    def run_grammar_checker(self):
        """
        Runs the grammar checker on the text and highlights issues.
        """
        text = self.text_edit.toPlainText()
        grammar_issues = check_grammar(text)
        self.highlight_issues(grammar_issues)
        self.display_recommendations(grammar_issues)

    def run_spell_checker(self):
        """
        Runs the spell checker on the text and highlights issues.
        """
        text = self.text_edit.toPlainText()
        spelling_issues = check_spelling(text)
        self.highlight_issues(spelling_issues)
        self.display_recommendations(spelling_issues)
        

    def highlight_issues(self, issues):
        """
        Highlights the issues found in the text.

        Args:
            issues (list): A list of dictionaries containing details about the issues.
        """
        highlighter = Highlighter(self.text_edit.document())
        highlighter.set_issues(issues)

    def display_recommendations(self, issues):
        """
        Displays the recommendations for the issues in the list widget.

        Args:
            issues (list): A list of dictionaries containing details about the issues.
        """
        self.recommendations_box.clear()
        for issue in issues:
            suggestions = ", ".join(issue['suggestions'])
            self.recommendations_box.addItem(f"{issue['message']} '{issue['word']}': {suggestions}")
        self.setProgressPercentage(issues)
        
    def setProgressPercentage(self,issues):
        text = self.text_edit.toPlainText()
        words = text.split()
        total_words = len(words) if len(words) != 0 else 1
        percentage = (len(issues) / total_words) * 100
        accuracy_percentage = round(100 - percentage,2)
        self.progress_bar.setValue(accuracy_percentage)
            
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())